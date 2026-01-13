"""Document parser for PDF, DOCX, Markdown, and text files."""

import io
import re
from pathlib import Path
from typing import Optional, Union

from ..config import logger, Settings
from ..models.enums import DocumentFormat
from ..models.policy_document import PolicyDocument, DocumentSection


class DocumentParser:
    """Parser for extracting text from various document formats."""

    SUPPORTED_EXTENSIONS = {
        ".pdf": DocumentFormat.PDF,
        ".docx": DocumentFormat.DOCX,
        ".doc": DocumentFormat.DOCX,
        ".md": DocumentFormat.MARKDOWN,
        ".markdown": DocumentFormat.MARKDOWN,
        ".txt": DocumentFormat.TEXT,
        ".text": DocumentFormat.TEXT,
    }

    def __init__(self, settings: Optional[Settings] = None):
        """Initialize document parser."""
        self.settings = settings or Settings()

    def parse_file(self, file_path: Union[str, Path]) -> PolicyDocument:
        """
        Parse a document file and extract its content.

        Args:
            file_path: Path to the document file

        Returns:
            PolicyDocument with extracted content

        Raises:
            ValueError: If file format is not supported or file is too large
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Check file size
        file_size = path.stat().st_size
        if file_size > self.settings.max_file_size_bytes:
            raise ValueError(
                f"File too large: {file_size / 1024 / 1024:.1f}MB "
                f"(max {self.settings.max_file_size_mb}MB)"
            )

        # Determine format
        doc_format = self._detect_format(path)
        if not doc_format:
            raise ValueError(f"Unsupported file format: {path.suffix}")

        # Parse based on format
        with open(path, "rb") as f:
            content = f.read()

        return self.parse_bytes(content, path.name, doc_format)

    def parse_bytes(
        self,
        content: bytes,
        filename: str,
        doc_format: Optional[DocumentFormat] = None
    ) -> PolicyDocument:
        """
        Parse document from bytes.

        Args:
            content: Document content as bytes
            filename: Original filename
            doc_format: Document format (auto-detected if not provided)

        Returns:
            PolicyDocument with extracted content
        """
        if not doc_format:
            doc_format = self._detect_format(Path(filename))

        if not doc_format:
            raise ValueError(f"Cannot determine format for: {filename}")

        # Validate size
        if len(content) > self.settings.max_file_size_bytes:
            raise ValueError(
                f"Content too large: {len(content) / 1024 / 1024:.1f}MB "
                f"(max {self.settings.max_file_size_mb}MB)"
            )

        # Parse by format
        if doc_format == DocumentFormat.PDF:
            text, sections, page_count = self._parse_pdf(content)
        elif doc_format == DocumentFormat.DOCX:
            text, sections, page_count = self._parse_docx(content)
        elif doc_format == DocumentFormat.MARKDOWN:
            text, sections, page_count = self._parse_markdown(content)
        else:  # TEXT
            text, sections, page_count = self._parse_text(content)

        # Extract title from content
        title = self._extract_title(text, filename)

        return PolicyDocument(
            filename=filename,
            format=doc_format,
            title=title,
            content=text,
            sections=sections,
            page_count=page_count,
            file_size_bytes=len(content),
        )

    def _detect_format(self, path: Path) -> Optional[DocumentFormat]:
        """Detect document format from file extension."""
        return self.SUPPORTED_EXTENSIONS.get(path.suffix.lower())

    def _parse_pdf(self, content: bytes) -> tuple[str, list[DocumentSection], int]:
        """Parse PDF document."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)

            sections = []
            all_text = []
            offset = 0

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                all_text.append(page_text)

                if page_text.strip():
                    sections.append(DocumentSection(
                        title=f"Page {i + 1}",
                        content=page_text,
                        page_number=i + 1,
                        start_offset=offset,
                        end_offset=offset + len(page_text),
                    ))
                offset += len(page_text) + 1

            text = "\n".join(all_text)
            return text, sections, page_count

        except Exception as e:
            logger.warning(f"PDF parsing error: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")

    def _parse_docx(self, content: bytes) -> tuple[str, list[DocumentSection], int]:
        """Parse DOCX document."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            sections = []
            all_text = []
            offset = 0
            current_section = None
            current_content = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if this is a heading
                if para.style.name.startswith("Heading"):
                    # Save previous section
                    if current_section and current_content:
                        section_text = "\n".join(current_content)
                        sections.append(DocumentSection(
                            title=current_section,
                            content=section_text,
                            start_offset=offset,
                            end_offset=offset + len(section_text),
                        ))
                        offset += len(section_text) + 1

                    current_section = text
                    current_content = []
                else:
                    current_content.append(text)
                    all_text.append(text)

            # Save last section
            if current_section and current_content:
                section_text = "\n".join(current_content)
                sections.append(DocumentSection(
                    title=current_section,
                    content=section_text,
                    start_offset=offset,
                    end_offset=offset + len(section_text),
                ))

            # If no sections found, create one
            if not sections and all_text:
                full_text = "\n".join(all_text)
                sections.append(DocumentSection(
                    title="Document",
                    content=full_text,
                    start_offset=0,
                    end_offset=len(full_text),
                ))

            text = "\n".join(all_text)
            return text, sections, 1  # DOCX doesn't have clear page breaks

        except Exception as e:
            logger.warning(f"DOCX parsing error: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")

    def _parse_markdown(self, content: bytes) -> tuple[str, list[DocumentSection], int]:
        """Parse Markdown document."""
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        sections = []
        offset = 0

        # Split by headings
        heading_pattern = r"^(#{1,6})\s+(.+)$"
        lines = text.split("\n")
        current_section = None
        current_content = []

        for line in lines:
            match = re.match(heading_pattern, line)
            if match:
                # Save previous section
                if current_section and current_content:
                    section_text = "\n".join(current_content)
                    sections.append(DocumentSection(
                        title=current_section,
                        content=section_text,
                        start_offset=offset,
                        end_offset=offset + len(section_text),
                    ))
                    offset += len(section_text) + 1

                current_section = match.group(2)
                current_content = []
            else:
                current_content.append(line)

        # Save last section
        if current_section and current_content:
            section_text = "\n".join(current_content)
            sections.append(DocumentSection(
                title=current_section,
                content=section_text,
                start_offset=offset,
                end_offset=offset + len(section_text),
            ))

        # If no sections, create one
        if not sections:
            sections.append(DocumentSection(
                title="Document",
                content=text,
                start_offset=0,
                end_offset=len(text),
            ))

        return text, sections, 1

    def _parse_text(self, content: bytes) -> tuple[str, list[DocumentSection], int]:
        """Parse plain text document."""
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        sections = [DocumentSection(
            title="Document",
            content=text,
            start_offset=0,
            end_offset=len(text),
        )]

        return text, sections, 1

    def _extract_title(self, text: str, filename: str) -> str:
        """Extract document title from content or filename."""
        # Try to find first heading
        lines = text.strip().split("\n")
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            # Markdown heading
            if line.startswith("#"):
                return re.sub(r"^#+\s*", "", line)
            # All caps title
            if line.isupper() and len(line) > 5:
                return line.title()
            # First non-empty line if reasonable length
            if line and 5 < len(line) < 100:
                return line

        # Fall back to filename
        return Path(filename).stem.replace("_", " ").replace("-", " ").title()

    def get_supported_formats(self) -> list[str]:
        """Get list of supported file extensions."""
        return list(self.SUPPORTED_EXTENSIONS.keys())


__all__ = ["DocumentParser"]
